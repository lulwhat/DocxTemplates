import docx
import re
import os

class SavePathIsNotAbsoluteError(Exception):
    pass
class SaveFileWrongFormatError(Exception):
    pass

class DocxHandler:
    def __init__(self, template_path):
        self.template_path = template_path
        self.doc = docx.Document(template_path)

    def docxSave(self, save_path):
        if "." not in os.path.basename(save_path):
            raise SaveFileWrongFormatError("Save file has no extension")
        elif (
            (os.path.basename(save_path).split(".")[1] != "doc") and
            (os.path.basename(save_path).split(".")[1] != "docx")
            ):
            raise SaveFileWrongFormatError("Save file extension should be .doc or .docx")
        elif os.path.isabs(save_path) == False:
            raise SavePathIsNotAbsoluteError("Need absolute file path")
        else:
            self.doc.save(save_path)
            return True

    def templateRead(self):
        vars = {}
        pars = self.doc.paragraphs
        for par in pars:
            find_vars = re.findall(r".*\{(.*)\}.*", par.text)
            for var in find_vars:
                vars[var] = ""
        return vars

    def docxReplace(self, template_path, data):
        self.doc = docx.Document(template_path)
        paragraphs = list(self.doc.paragraphs)
        # for t in self.doc.tables:
        #     for row in t.rows:
        #         for cell in row.cells:
        #             for paragraph in cell.paragraphs:
        #                 paragraphs.append(paragraph)
        for sec in self.doc.sections:
            for paragraph in sec.footer.paragraphs:
                paragraphs.append(paragraph)
            for paragraph in sec.header.paragraphs:
                paragraphs.append(paragraph)
        for p in paragraphs:
            for key, val in data.items():
                key_name = '{{{}}}'.format(key) # I'm using placeholders in the form {PlaceholderName}
                if key_name in p.text:
                    inline = p.runs
                    # Replace strings and retain the same style.
                    # The text to be replaced can be split over several runs so
                    # search through, identify which runs need to have text replaced
                    # then replace the text in those identified
                    started = False
                    key_index = 0
                    # found_runs is a list of (inline index, index of match, length of match)
                    found_runs = list()
                    found_all = False
                    replace_done = False
                    for i in range(len(inline)):

                        # case 1: found in single run so short circuit the replace
                        if key_name in inline[i].text and not started:
                            found_runs.append((i, inline[i].text.find(key_name), len(key_name)))
                            text = inline[i].text.replace(key_name, str(val))
                            inline[i].text = text
                            replace_done = True
                            found_all = True
                            break

                        if key_name[key_index] not in inline[i].text and not started:
                            # keep looking ...
                            continue

                        # case 2: search for partial text, find first run
                        if key_name[key_index] in inline[i].text and inline[i].text[-1] in key_name and not started:
                            # check sequence
                            start_index = inline[i].text.find(key_name[key_index])
                            check_length = len(inline[i].text)
                            for text_index in range(start_index, check_length):
                                if inline[i].text[text_index] != key_name[key_index]:
                                    # no match so must be false positive
                                    break
                            if key_index == 0:
                                started = True
                            chars_found = check_length - start_index
                            key_index += chars_found
                            found_runs.append((i, start_index, chars_found))
                            if key_index != len(key_name):
                                continue
                            else:
                                # found all chars in key_name
                                found_all = True
                                break

                        # case 2: search for partial text, find subsequent run
                        if key_name[key_index] in inline[i].text and started and not found_all:
                            # check sequence
                            chars_found = 0
                            check_length = len(inline[i].text)
                            for text_index in range(0, check_length):
                                if inline[i].text[text_index] == key_name[key_index]:
                                    key_index += 1
                                    chars_found += 1
                                else:
                                    break
                            # no match so must be end
                            found_runs.append((i, 0, chars_found))
                            if key_index == len(key_name):
                                found_all = True
                                break

                    if found_all and not replace_done:
                        for i, item in enumerate(found_runs):
                            index, start, length = [t for t in item]
                            if i == 0:
                                text = inline[index].text.replace(inline[index].text[start:start + length], str(val))
                                inline[index].text = text
                            else:
                                text = inline[index].text.replace(inline[index].text[start:start + length], '')
                                inline[index].text = text