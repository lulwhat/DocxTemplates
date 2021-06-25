import docx
import re
import os

class SavePathIsNotAbsoluteError(Exception):
    pass
class SaveFileWrongFormatError(Exception):
    pass

class DocxHandler:
    def __init__(self, template_path):
        self.template_path = template_path
        self.doc = docx.Document(template_path)

    def docxSave(self, save_path):
        if "." not in os.path.basename(save_path):
            raise SaveFileWrongFormatError("Save file has no extension")
        elif (
            (os.path.basename(save_path).split(".")[1] != "doc") and
            (os.path.basename(save_path).split(".")[1] != "docx")
            ):
            raise SaveFileWrongFormatError("Save file extension should be .doc or .docx")
        elif os.path.isabs(save_path) == False:
            raise SavePathIsNotAbsoluteError("Need absolute file path")
        else:
            self.doc.save(save_path)
            return True

    def templateRead(self):
        vars = {}
        pars = list(self.doc.paragraphs)
        secs = list(self.doc.sections)
        tabs = list(self.doc.tables)

        # add footers and headers pars
        for sec in secs:
            for par in sec.footer.paragraphs:
                pars.append(par)
            for par in sec.header.paragraphs:
                pars.append(par)
        
        # add tabs pars
        for tab in tabs:
            for row in tab.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        pars.append(par)

        for par in pars:
            find_vars = re.findall(r".*\{(.*)\}.*", par.text)
            for var in find_vars:
                vars[var] = ""
        return vars

    def docxReplace(self, template_path, data):
        self.doc = docx.Document(template_path)
        pars = list(self.doc.paragraphs)
        secs = list(self.doc.sections)
        tabs = list(self.doc.tables)

        # add footers and headers pars
        for sec in secs:
            for par in sec.footer.paragraphs:
                pars.append(par)
            for par in sec.header.paragraphs:
                pars.append(par)
                
        # add tabs pars
        for tab in tabs:
            for row in tab.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        pars.append(par)

        # parse paragraphs and find variables to replace
        for par in pars:
            par_vars = re.findall(r".*\{(.*)\}.*", par.text)
            if len(par_vars) > 0:
                for var in par_vars:
                    self.replaceFound(par, var, data)

    def replaceFound(self, par, var, data):
        for i in range(len(par.runs)):
            # replace "{var}" and "var" cases
            if par.runs[i].text == f"{var}":
                par.runs[i].text = par.runs[i].text.replace(f"{var}", data[var])
            elif par.runs[i].text == f"{{{var}}}":
                par.runs[i].text = par.runs[i].text.replace(f"{{{var}}}", data[var])
        # clear curly braces after replacements
        for i in range(len(par.runs)):
            for rep in (("{", ""), ("}", "")):
                par.runs[i].text = par.runs[i].text.replace(*rep)